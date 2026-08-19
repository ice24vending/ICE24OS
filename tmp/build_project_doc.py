from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ICE24_OS_Documentacion_Actual_v1.0.docx"
ASSETS = ROOT / "tmp" / "doc_assets"

NAVY = "04144C"
ACTION = "0455C4"
SKY = "9CD5EF"
PALE = "EAF5FB"
WHITE = "FFFFFF"
INK = "172033"
MUTED = "5E687A"
LINE = "D7DEE8"
SOFT = "F2F4F7"
SUCCESS = "16794D"
WARNING = "A15C00"
CRITICAL = "B42318"
INFO = "1769AA"
OFFLINE = "007B83"


def rgb(hex_color: str) -> tuple[int, int, int]:
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/seguisb.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rr(draw: ImageDraw.ImageDraw, xy, radius: int, fill: str, outline: str | None = None, width: int = 1):
    draw.rounded_rectangle(xy, radius=radius, fill=rgb(fill), outline=rgb(outline) if outline else None, width=width)


def desktop_mockup(path: Path) -> None:
    im = Image.new("RGB", (1600, 980), rgb("E9EEF5"))
    d = ImageDraw.Draw(im)
    # Browser frame
    rr(d, (28, 28, 1572, 952), 28, WHITE, "C5CEDA", 3)
    d.rectangle((28, 28, 1572, 100), fill=rgb("F7F9FC"))
    for x, c in [(65, "ED6A5E"), (98, "F4BF4F"), (131, "61C454")]:
        d.ellipse((x-10, 55, x+10, 75), fill=rgb(c))
    rr(d, (390, 48, 1210, 84), 16, WHITE, "D7DEE8")
    d.text((645, 55), "app.ice24.mx / mantenimiento", font=font(17), fill=rgb(MUTED))

    # Sidebar
    d.rectangle((28, 100, 335, 952), fill=rgb(NAVY))
    d.text((70, 142), "ICE24", font=font(40, True), fill=rgb(WHITE))
    d.text((211, 153), "OS", font=font(21, True), fill=rgb(SKY))
    d.text((70, 205), "Cuenta piloto · Norte", font=font(16), fill=rgb("C5D8F8"))
    nav = [("Resumen", True), ("Órdenes", False), ("Máquinas", False), ("Sucursales", False), ("Historial", False), ("Auditoría", False)]
    y = 276
    for label, active in nav:
        if active:
            rr(d, (52, y-12, 311, y+40), 12, ACTION)
        d.text((82, y), label, font=font(21, active), fill=rgb(WHITE if active else "C8D2E4"))
        y += 74
    rr(d, (54, 832, 306, 910), 14, "092161", "27427F")
    d.ellipse((76, 852, 112, 888), fill=rgb(SKY))
    d.text((128, 849), "Ana Torres", font=font(18, True), fill=rgb(WHITE))
    d.text((128, 876), "Propietaria", font=font(15), fill=rgb("C8D2E4"))

    # Main header
    d.text((385, 145), "Resumen de mantenimiento", font=font(38, True), fill=rgb(INK))
    d.text((386, 198), "Martes 18 de agosto · Sucursal Centro", font=font(18), fill=rgb(MUTED))
    rr(d, (1280, 140, 1518, 202), 15, ACTION)
    d.text((1325, 157), "+ Nueva orden", font=font(20, True), fill=rgb(WHITE))

    # KPI cards
    cards = [
        ("Órdenes abiertas", "12", ACTION, "+3 esta semana"),
        ("Por vencer", "4", WARNING, "Próximas 48 h"),
        ("Completadas", "37", SUCCESS, "92% a tiempo"),
    ]
    x = 385
    for title, value, color, note in cards:
        rr(d, (x, 250, x+350, 425), 18, WHITE, LINE, 2)
        d.text((x+28, 280), title, font=font(18), fill=rgb(MUTED))
        d.text((x+28, 322), value, font=font(44, True), fill=rgb(color))
        d.text((x+102, 344), note, font=font(16), fill=rgb(MUTED))
        x += 380

    # Orders panel
    rr(d, (385, 468, 1518, 890), 18, WHITE, LINE, 2)
    d.text((415, 498), "Trabajo prioritario", font=font(23, True), fill=rgb(INK))
    d.text((1338, 502), "Ver todas →", font=font(17, True), fill=rgb(ACTION))
    headers = [(415, "Orden"), (650, "Máquina"), (890, "Responsable"), (1130, "Vence"), (1320, "Estado")]
    d.rectangle((405, 548, 1498, 600), fill=rgb(SOFT))
    for hx, label in headers:
        d.text((hx, 563), label, font=font(16, True), fill=rgb(MUTED))
    rows = [
        ("OT-000184", "ICE-024", "Luis García", "Hoy 16:00", "En curso", INFO),
        ("OT-000185", "ICE-011", "Marta Díaz", "Mañana", "Asignada", ACTION),
        ("OT-000179", "ICE-007", "Luis García", "Vencida", "Atención", CRITICAL),
        ("OT-000186", "ICE-031", "Sin asignar", "20 ago", "Borrador", MUTED),
    ]
    y = 624
    for idx, (order, machine, owner, due, status, color) in enumerate(rows):
        if idx:
            d.line((415, y-17, 1488, y-17), fill=rgb("E7EBF0"), width=2)
        d.text((415, y), order, font=font(17, True), fill=rgb(ACTION))
        d.text((650, y), machine, font=font(17), fill=rgb(INK))
        d.text((890, y), owner, font=font(17), fill=rgb(INK))
        d.text((1130, y), due, font=font(17), fill=rgb(CRITICAL if due == "Vencida" else INK))
        rr(d, (1315, y-7, 1458, y+34), 14, "EAF3FF" if color == ACTION else ("E8F6EF" if color == SUCCESS else ("FDECEC" if color == CRITICAL else "E8F2FA")))
        d.text((1333, y+2), status, font=font(15, True), fill=rgb(color))
        y += 67
    im.save(path, quality=94)


def mobile_mockup(path: Path) -> None:
    im = Image.new("RGB", (900, 1500), rgb("E9EEF5"))
    d = ImageDraw.Draw(im)
    # Phone shell
    rr(d, (160, 35, 740, 1465), 70, "172033")
    rr(d, (184, 70, 716, 1428), 48, WHITE)
    rr(d, (337, 53, 565, 86), 16, "172033")
    # App header
    d.rectangle((184, 70, 716, 250), fill=rgb(NAVY))
    d.text((226, 118), "‹", font=font(40, True), fill=rgb(WHITE))
    d.text((285, 115), "Orden OT-000184", font=font(25, True), fill=rgb(WHITE))
    rr(d, (522, 120, 680, 169), 18, "0B3077")
    d.text((548, 132), "● En curso", font=font(16, True), fill=rgb(SKY))
    d.text((226, 190), "Máquina ICE-024 · Sucursal Centro", font=font(17), fill=rgb("C8D8F0"))

    # Offline banner
    rr(d, (208, 275, 692, 345), 16, "E6F6F6", "B6DFDF", 2)
    d.text((235, 295), "SIN RED · guardando en el dispositivo", font=font(16, True), fill=rgb(OFFLINE))

    d.text((208, 382), "Checklist preventivo", font=font(27, True), fill=rgb(INK))
    d.text((208, 425), "3 de 5 actividades completadas", font=font(17), fill=rgb(MUTED))
    d.rectangle((208, 468, 692, 481), fill=rgb("DFE5EC"))
    d.rectangle((208, 468, 502, 481), fill=rgb(SUCCESS))

    items = [
        ("Inspección visual", True, "Sin daños visibles"),
        ("Limpieza de condensador", True, "Completada"),
        ("Verificar temperatura", True, "3.4 °C"),
        ("Revisar conexiones", False, "Pendiente"),
        ("Fotografía final", False, "Requerida"),
    ]
    y = 525
    for title, done, note in items:
        rr(d, (208, y, 692, y+118), 16, WHITE, LINE, 2)
        if done:
            d.ellipse((230, y+32, 276, y+78), fill=rgb(SUCCESS))
            d.line((241, y+55, 250, y+65), fill=rgb(WHITE), width=4)
            d.line((250, y+65, 266, y+45), fill=rgb(WHITE), width=4)
        else:
            d.ellipse((230, y+32, 276, y+78), fill=rgb(WHITE), outline=rgb("9AA6B4"), width=3)
        d.text((297, y+26), title, font=font(18, True), fill=rgb(INK))
        d.text((297, y+61), note, font=font(16), fill=rgb(MUTED))
        y += 134

    rr(d, (208, 1220, 692, 1290), 18, PALE, ACTION, 2)
    d.text((323, 1241), "+ Agregar evidencia", font=font(19, True), fill=rgb(ACTION))
    rr(d, (208, 1310, 692, 1384), 18, ACTION)
    d.text((300, 1333), "Guardar avance local", font=font(20, True), fill=rgb(WHITE))
    im.save(path, quality=94)


def architecture_diagram(path: Path) -> None:
    im = Image.new("RGB", (1600, 920), rgb(WHITE))
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Arquitectura objetivo del piloto en Google Cloud", font=font(34, True), fill=rgb(NAVY))
    d.text((70, 93), "Servicios administrados, contenedores portables y datos en la región de México", font=font(19), fill=rgb(MUTED))
    # users
    rr(d, (70, 250, 310, 410), 20, PALE, SKY, 3)
    d.text((135, 285), "Usuarios", font=font(25, True), fill=rgb(NAVY))
    d.text((105, 330), "PWA · escritorio", font=font(18), fill=rgb(INK))
    d.text((118, 360), "móvil · portal", font=font(18), fill=rgb(INK))
    # edge
    rr(d, (395, 210, 685, 450), 20, NAVY)
    d.text((463, 248), "Borde web", font=font(25, True), fill=rgb(WHITE))
    d.text((433, 300), "HTTPS Load Balancing", font=font(18), fill=rgb("D5E2FA"))
    d.text((485, 335), "Cloud CDN", font=font(18), fill=rgb("D5E2FA"))
    d.text((475, 370), "Cloud Armor", font=font(18), fill=rgb("D5E2FA"))
    # cloud run services
    rr(d, (770, 150, 1125, 510), 24, "F3F7FC", "B9CAE4", 3)
    d.text((850, 180), "Cloud Run", font=font(28, True), fill=rgb(ACTION))
    services = ["PWA + BFF", "API NestJS", "Portal público", "Keycloak (mín. 1)"]
    y = 240
    for service in services:
        rr(d, (815, y, 1080, y+52), 12, WHITE, LINE, 2)
        d.text((850, y+14), service, font=font(18, True), fill=rgb(INK))
        y += 67
    # data
    data_boxes = [
        (1200, 140, 1515, 255, "Cloud SQL", "PostgreSQL + PostGIS"),
        (1200, 285, 1515, 400, "Cloud Storage", "cuarentena · originales"),
        (1200, 430, 1515, 545, "Secret Manager", "KMS · secretos"),
    ]
    for x1, y1, x2, y2, title, sub in data_boxes:
        rr(d, (x1, y1, x2, y2), 18, PALE, SKY, 2)
        d.text((x1+28, y1+22), title, font=font(21, True), fill=rgb(NAVY))
        d.text((x1+28, y1+64), sub, font=font(16), fill=rgb(MUTED))
    # async and ops
    rr(d, (395, 625, 760, 810), 20, "F7F9FC", LINE, 2)
    d.text((448, 655), "Procesamiento asíncrono", font=font(22, True), fill=rgb(NAVY))
    d.text((443, 700), "Pub/Sub · Cloud Tasks", font=font(18), fill=rgb(INK))
    d.text((443, 735), "Scheduler · Jobs · PDF", font=font(18), fill=rgb(INK))
    d.text((443, 770), "ClamAV por Eventarc", font=font(18), fill=rgb(INK))
    rr(d, (830, 625, 1195, 810), 20, "F7F9FC", LINE, 2)
    d.text((930, 655), "Operación", font=font(22, True), fill=rgb(NAVY))
    d.text((874, 700), "Artifact Registry · Terraform", font=font(18), fill=rgb(INK))
    d.text((881, 735), "Logging · Monitoring · Trace", font=font(18), fill=rgb(INK))
    d.text((923, 770), "OpenTelemetry", font=font(18), fill=rgb(INK))
    # arrows
    for (x1, y1, x2, y2) in [(310, 330, 395, 330), (685, 330, 770, 330), (1125, 210, 1200, 200), (1125, 335, 1200, 340), (1125, 450, 1200, 485), (950, 510, 700, 625), (1000, 510, 1000, 625)]:
        d.line((x1, y1, x2, y2), fill=rgb(ACTION), width=5)
        d.polygon([(x2, y2), (x2-14, y2-8), (x2-14, y2+8)], fill=rgb(ACTION))
    im.save(path, quality=94)


def flow_diagram(path: Path) -> None:
    im = Image.new("RGB", (1600, 520), rgb(WHITE))
    d = ImageDraw.Draw(im)
    d.text((70, 40), "Escenario operativo del MVP-1", font=font(32, True), fill=rgb(NAVY))
    steps = [
        ("1", "Acceso", "OIDC + 2FA"),
        ("2", "Contexto", "cuenta/sucursal"),
        ("3", "Orden", "crear y asignar"),
        ("4", "Campo", "checklist/fotos"),
        ("5", "Offline", "guardar local"),
        ("6", "Sincronizar", "resolver conflicto"),
        ("7", "Cierre", "historial/auditoría"),
    ]
    x = 70
    for i, (num, title, sub) in enumerate(steps):
        rr(d, (x, 150, x+185, 360), 20, PALE if i % 2 == 0 else "F4F7FC", SKY if i % 2 == 0 else LINE, 2)
        d.ellipse((x+62, 170, x+122, 230), fill=rgb(ACTION))
        d.text((x+82, 182), num, font=font(23, True), fill=rgb(WHITE))
        tw = d.textlength(title, font=font(19, True))
        d.text((x+(185-tw)/2, 258), title, font=font(19, True), fill=rgb(NAVY))
        sw = d.textlength(sub, font=font(15))
        d.text((x+(185-sw)/2, 302), sub, font=font(15), fill=rgb(MUTED))
        if i < len(steps)-1:
            d.line((x+185, 255, x+215, 255), fill=rgb(ACTION), width=4)
            d.polygon([(x+215, 255), (x+202, 247), (x+202, 263)], fill=rgb(ACTION))
        x += 215
    d.text((70, 420), "La evidencia capturada sin red conserva trazabilidad y sólo se publica después de sincronización y validación del servidor.", font=font(18), fill=rgb(MUTED))
    im.save(path, quality=94)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    tr_pr.append(el)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    # left/right are more interoperable with Word's PDF exporter than start/end.
    for m, v in [("top", top), ("left", start), ("bottom", bottom), ("right", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths: Sequence[int], header=True) -> None:
    # Reserve 120 dxa on each side inside the 9360-dxa text area. This avoids
    # a Word PDF-export edge case where the first cell of a full-bleed table
    # can be positioned outside its border.
    target_width = 9120
    total = sum(widths)
    widths = [round(w * target_width / total) for w in widths]
    widths[-1] += target_width - sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(target_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    for row_idx, row in enumerate(table.rows):
        prevent_row_split(row)
        if header and row_idx == 0:
            set_repeat_table_header(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.right_indent = Inches(0)
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
            if header and row_idx == 0:
                set_cell_shading(cell, SOFT)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    table.style = "Table Grid"


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int]):
    data = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in data:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = str(value)
    configure_table(table, widths, header=True)
    return table


def set_alt_text(picture_run, descr: str) -> None:
    for doc_pr in picture_run._r.xpath(".//wp:docPr"):
        doc_pr.set("descr", descr)


def add_picture(doc: Document, path: Path, width=Inches(6.35), alt=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    pic = run.add_picture(str(path), width=width)
    if alt:
        set_alt_text(run, alt)
    return pic


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_callout(doc: Document, title: str, body: str, color=ACTION, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(body)
    configure_table(table, [9360], header=False)
    # left accent
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left = OxmlElement("w:start")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), color)
    borders.append(left)


def add_bullet(doc: Document, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_numbered_list(doc: Document, items: Sequence[str]):
    """Add a real numbered list that restarts at 1 for this logical group."""
    numbering = doc.part.numbering_part.element
    base_num_id = int(doc.styles["List Number"].element.pPr.numPr.numId.val)
    base_num = numbering.num_having_numId(base_num_id)
    num = numbering.add_num(base_num.abstractNumId.val)
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    num_id = int(num.numId)
    paragraphs = []
    for item in items:
        para = doc.add_paragraph(style="List Number")
        num_pr = para._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
        para.add_run(item)
        paragraphs.append(para)
    return paragraphs


def field(run, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def page_break(doc: Document):
    # Section starts are controlled by Heading 1/page_break_before. Keeping the
    # break on the heading avoids accidental blank pages when the prior page is full.
    return None


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Title": (34, NAVY, 0, 12),
        "Subtitle": (16, ACTION, 0, 12),
        "Heading 1": (16, ACTION, 16, 8),
        "Heading 2": (13, ACTION, 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            st.paragraph_format.page_break_before = True
    styles["Caption"].font.name = "Calibri"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].paragraph_format.space_after = Pt(8)
    for list_name in ["List Bullet", "List Bullet 2", "List Number"]:
        st = styles[list_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5 if list_name != "List Bullet 2" else 0.75)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167

    # Running header/footer
    header = sec.header
    hp = header.paragraphs[0]
    hp.text = "ICE24 OS  |  Documentación maestra del proyecto"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Versión 1.0 · 18/08/2026  |  Página ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    pg = fp.add_run()
    field(pg, "PAGE")
    fp.add_run(" de ")
    total = fp.add_run()
    field(total, "NUMPAGES")

    doc.core_properties.title = "ICE24 OS — Documentación actual del proyecto"
    doc.core_properties.subject = "Estado actual, arquitectura, alcance y visión funcional del MVP-1"
    doc.core_properties.author = "Equipo ICE24 OS"
    doc.core_properties.keywords = "ICE24 OS, PWA, mantenimiento, Google Cloud, MVP"
    doc.core_properties.comments = "Documento generado a partir de la documentación versionada del repositorio."


def p(doc: Document, text: str = "", bold_lead: str | None = None):
    para = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        para.add_run(bold_lead).bold = True
        para.add_run(text[len(bold_lead):])
    else:
        para.add_run(text)
    return para


def build() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    desktop = ASSETS / "mvp_desktop.png"
    mobile = ASSETS / "mvp_mobile.png"
    architecture = ASSETS / "architecture.png"
    flow = ASSETS / "mvp_flow.png"
    desktop_mockup(desktop)
    mobile_mockup(mobile)
    architecture_diagram(architecture)
    flow_diagram(flow)

    doc = Document()
    configure_document(doc)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = label.add_run("ICE24")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = label.add_run(" OS")
    r2.bold = True
    r2.font.size = Pt(17)
    r2.font.color.rgb = RGBColor.from_string(ACTION)
    title = doc.add_paragraph(style="Title")
    title.add_run("Documentación actual\ndel proyecto")
    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run("Estado funcional y técnico · definición del MVP-1 · visión de la primera versión")
    doc.add_paragraph()
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(24)
    ppr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "20")
    bottom.set(qn("w:color"), SKY)
    pbdr.append(bottom)
    ppr.append(pbdr)
    meta = doc.add_paragraph()
    meta.add_run("Versión documental\n").bold = True
    meta.add_run("1.0 — línea base posterior al avance de Fase 0\n\n")
    meta.add_run("Fecha de corte\n").bold = True
    meta.add_run("18 de agosto de 2026\n\n")
    meta.add_run("Estado\n").bold = True
    meta.add_run("Borrador consolidado para revisión y aprobación\n\n")
    meta.add_run("Plataforma objetivo\n").bold = True
    meta.add_run("Aplicación web progresiva desplegada en Google Cloud, región México")
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Documento confidencial de trabajo · ICE24").italic = True
    note.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    page_break(doc)
    doc.add_heading("Control del documento", level=1)
    add_table(doc, ["Campo", "Valor"], [
        ("Propósito", "Consolidar el estado actual del proyecto y servir como referencia para Fase 1."),
        ("Fuentes", "context/*.md y docs/*.md del repositorio ICE24OS."),
        ("Audiencia", "Dirección, Producto, Ingeniería, Seguridad, Operación, Sanidad y Jurídico."),
        ("Alcance", "Decisiones vigentes, requisitos base, MVP-1, arquitectura, riesgos y ruta inmediata."),
        ("No sustituye", "Aprobaciones de presupuesto, privacidad, sanidad, contratos ni validación con usuarios."),
    ], [1800, 7560])
    doc.add_heading("Resumen ejecutivo", level=1)
    p(doc, "ICE24 OS se encuentra en una fase de definición avanzada: la documentación funcional y técnica ya permite iniciar el monorepo, desplegar una primera URL de desarrollo y construir una rebanada vertical de mantenimiento sin utilizar datos productivos.")
    add_callout(doc, "Decisión principal", "La primera entrega real será el MVP-1: una PWA privada para una cuenta piloto que permite gestionar sucursales y máquinas, crear y asignar órdenes de mantenimiento, ejecutar checklists con fotografías incluso con conectividad intermitente, sincronizar, revisar historial y conservar auditoría.")
    add_table(doc, ["Dimensión", "Línea base actual"], [
        ("Nube", "Google Cloud · northamerica-south1 (México)"),
        ("Aplicación", "PWA responsive · Next.js + NestJS · TypeScript"),
        ("Datos", "Cloud SQL PostgreSQL/PostGIS + Cloud Storage privado"),
        ("Identidad", "Keycloak · OIDC Authorization Code + PKCE · patrón BFF"),
        ("Capacidad piloto", "1 cuenta · hasta 5 sucursales · 25 máquinas · 20 usuarios (supuesto)"),
        ("Costo objetivo", "USD 100–300/mes; tope provisional de control: USD 400 producción"),
        ("Disponibilidad", "99.5% mensual para PWA+BFF y API privada"),
        ("Continuidad", "RPO 15 min · RTO 4 h para datos transaccionales"),
    ], [2100, 7260])

    page_break(doc)
    doc.add_heading("Contenido", level=1)
    contents = [
        "Visión, problema y objetivos", "Estado actual del proyecto", "Usuarios, roles y autorización",
        "Alcance funcional y releases", "MVP-1: definición de la primera versión", "Así se verá la primera versión",
        "Escenario operativo de punta a punta", "Arquitectura de solución en Google Cloud", "Stack y estructura técnica",
        "Datos, seguridad e identidad", "Offline, archivos, observabilidad y continuidad", "Entornos, costos y despliegue",
        "Calidad y criterios de aceptación", "Estado de Fase 0 y gates", "Riesgos, bloqueos y decisiones abiertas",
        "Ruta recomendada", "Anexos: decisiones y documentos fuente"
    ]
    add_numbered_list(doc, contents)
    add_callout(doc, "Cómo leer este documento", "Los elementos marcados como línea base o aceptados habilitan construcción. Los elementos en revisión requieren aprobación. Los bloqueos externos impiden activar funciones específicas, pero no detienen el esqueleto web ni el flujo piloto de mantenimiento.", INFO, "EEF6FB")

    page_break(doc)
    doc.add_heading("1. Visión, problema y objetivos", level=1)
    doc.add_heading("1.1 Visión", level=2)
    p(doc, "ICE24 OS será la plataforma web de operación y trazabilidad para administrar cuentas, sucursales, máquinas, mantenimiento, evidencias, cumplimiento y, en releases posteriores, inventario, ventas, pedidos y publicación técnica/sanitaria.")
    doc.add_heading("1.2 Problema que resuelve", level=2)
    for item in [
        "Información operativa dispersa entre hojas, mensajes, fotografías y archivos sin historial consistente.",
        "Dificultad para conocer qué máquina requiere atención, quién es responsable y qué evidencia existe.",
        "Riesgo de mezclar información entre empresas, sucursales o perfiles con distintos privilegios.",
        "Conectividad variable en campo y necesidad de conservar avance sin duplicar operaciones.",
        "Necesidad de publicación deliberada y separada de información privada, técnica o sanitaria.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("1.3 Objetivos del primer ciclo", level=2)
    add_table(doc, ["Objetivo", "Resultado observable"], [
        ("Disponibilidad web", "Entorno de desarrollo accesible por URL desde Fase 1."),
        ("Flujo completo", "Orden de mantenimiento creada, ejecutada, sincronizada y auditada."),
        ("Seguridad por diseño", "Separación por cuenta, autorización de servidor y mínimo privilegio."),
        ("Aprendizaje del piloto", "Medir adopción, cumplimiento, sincronización e incidentes durante 60 días."),
        ("Base extensible", "Monolito modular y contratos estables para agregar dominios sin rehacer la plataforma."),
    ], [2700, 6660])
    doc.add_heading("1.4 Principios", level=2)
    for item in [
        "Denegación por defecto; separación estricta entre identidad, cuenta, ámbito y acción.",
        "El servidor autoriza; ocultar la interfaz no concede ni revoca permisos.",
        "Una rebanada vertical útil antes de ampliar dominios.",
        "Infraestructura reproducible, imágenes inmutables y secretos fuera del repositorio.",
        "Operación offline explícita, auditable e idempotente.",
        "No publicar/eliminar información regulada sin aprobación y trazabilidad.",
    ]:
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("2. Estado actual del proyecto", level=1)
    p(doc, "La Fase 0 cuenta con entregables documentales para gobierno, producto, plataforma, stack, identidad, autorización, continuidad, integraciones, soporte, incidentes y backlog. El proyecto todavía no es una aplicación desplegada: la siguiente fase crea el monorepo, los servicios base y el primer entorno web.")
    add_table(doc, ["Área", "Estado", "Qué existe hoy"], [
        ("Producto", "En revisión", "MVP-1 definido, métricas de éxito y alcance excluido."),
        ("Arquitectura", "Línea base", "Google Cloud México, Cloud Run, Cloud SQL, Storage y mensajería."),
        ("Stack", "Aceptado", "Node 24 LTS, Next.js, NestJS, Vitest, Playwright y Terraform."),
        ("Identidad", "En revisión", "Keycloak, BFF, sesiones, 2FA y recuperación propuestos."),
        ("Autorización", "En revisión", "9 roles base, RBAC+ABAC y pruebas negativas obligatorias."),
        ("Continuidad", "En revisión", "SLO, RPO/RTO, retención y pruebas de restauración."),
        ("Backlog", "Documentado", "Épicas trazables y orden inicial de releases."),
        ("Implementación", "Pendiente Fase 1", "No hay todavía código de aplicación ni infraestructura aprovisionada."),
    ], [2100, 1900, 5360])
    add_callout(doc, "Gate habilitado", "Puede iniciarse Fase 1 y desplegarse un esqueleto con datos sintéticos. No se autoriza todavía cargar datos reales, publicar contenido sanitario o activar importaciones Excel.", SUCCESS, "EAF6F0")
    doc.add_heading("2.1 Madurez de las decisiones", level=2)
    for item in [
        "Aceptado para construcción: stack técnico y Vitest.",
        "Línea base pendiente de aprobación presupuestal/privacidad: plataforma Google Cloud.",
        "En revisión funcional u operativa: MVP, RACI, identidad, autorización, continuidad, proveedores, soporte e incidentes.",
        "Bloqueado por insumos externos: Excel reales, plantillas aprobadas y dictamen sanitario/jurídico.",
    ]:
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("3. Usuarios, roles y autorización", level=1)
    p(doc, "El modelo distingue identidad global, relación con una cuenta, ámbito de operación y acción solicitada. Cada comando es autorizado por la API mediante RBAC+ABAC; la regla base es denegar.")
    add_table(doc, ["Código", "Rol", "Ámbito máximo"], [
        ("IA", "ICE24 Admin", "Global, sólo acciones explícitas"),
        ("IO", "ICE24 Operaciones", "Global operativo, sin secretos ni cobros"),
        ("OW", "Propietario de cuenta", "Su cuenta y recursos asociados"),
        ("TC", "Técnico", "Asignaciones y máquinas autorizadas"),
        ("OP", "Operador", "Sucursal/máquinas asignadas"),
        ("SA", "Responsable sanitario", "Ámbito sanitario asignado"),
        ("DV", "Repartidor", "Asociaciones, pedidos y evidencia propios"),
        ("RA", "Restaurante Admin", "Su negocio y pedidos"),
        ("AU", "Consulta/Auditor", "Lectura/exportación expresamente concedida"),
    ], [900, 2500, 5960])
    doc.add_heading("3.1 Controles reservados", level=2)
    for item in [
        "Alta global, revocación global y overrides documentados: sólo ICE24 Admin.",
        "Publicación sanitaria: Responsable Sanitario y, cuando aplique, segunda aprobación ICE24.",
        "Recuperación manual: dos operadores distintos; nadie se autoaprueba.",
        "Cuenta en modo lectura: permite consulta autorizada y bloquea comandos.",
        "Descargas confidenciales y acciones críticas: auditadas y con reautenticación cuando corresponda.",
    ]:
        add_bullet(doc, item)
    heading = doc.add_heading("3.2 Clases de datos", level=2)
    heading.paragraph_format.page_break_before = True
    add_table(doc, ["Clase", "Ejemplos", "Regla"], [
        ("Pública", "Versión deliberadamente publicada", "Sólo datos aprobados para portal público"),
        ("Interna", "Estados y catálogos", "Usuario autenticado con módulo y ámbito"),
        ("Confidencial", "Contacto, ubicación, costos, ventas", "Mínimo privilegio; descarga auditada"),
        ("Restringida", "Recuperación, fiscal, originales, auditoría", "Acción nominativa; nunca en logs"),
    ], [1500, 3100, 4760])

    page_break(doc)
    doc.add_heading("4. Alcance funcional y releases", level=1)
    p(doc, "El alcance completo del producto permanece como roadmap. La estrategia evita intentar todos los dominios antes de obtener evidencia real con mantenimiento.")
    add_table(doc, ["Release", "Propósito", "Dominios principales"], [
        ("MVP-1", "Piloto web de mantenimiento", "Identidad, cuentas, sucursales, máquinas, órdenes, checklist, fotos, offline, historial y auditoría"),
        ("MVP-2", "Piloto técnico-sanitario", "Bitácoras, laboratorio/restricciones mínimas y documentos/portal indispensables, sólo tras validación"),
        ("Siguientes", "Expansión operativa", "Inventario, costos, ventas, tarjetas, pedidos, reparto, suscripción y publicación ampliada"),
    ], [1300, 2500, 5560])
    doc.add_heading("4.1 Incluido en MVP-1", level=2)
    for item in [
        "Inicio de sesión, recuperación y selección de contexto de cuenta.",
        "Administración mínima de una cuenta piloto, sucursales, máquinas y asociaciones de usuarios.",
        "Mantenimiento preventivo/correctivo, tickets, órdenes, asignación, prioridades y vencimientos.",
        "Ejecución móvil con checklist, notas, fotografías, guardado local y sincronización.",
        "Conflicto controlado, historial, auditoría y notificaciones transaccionales.",
        "PWA instalable y responsive para escritorio, Android e iOS soportados.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("4.2 Fuera de MVP-1", level=2)
    for item in [
        "Reglas sanitarias publicables, laboratorio y límites normativos.",
        "Portal público/QR y publicación de documentos regulados.",
        "Inventario avanzado, Excel de ventas, tarjetas, pedidos y reparto.",
        "Cobro automático; Stripe puede evaluarse en sandbox sin condicionar el piloto.",
        "Migración masiva; sólo carga inicial controlada y auditable.",
    ]:
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("5. MVP-1: definición de la primera versión", level=1)
    add_callout(doc, "Hipótesis de valor", "Si una cuenta piloto puede programar, ejecutar y demostrar mantenimiento desde una misma aplicación web —aun con red intermitente— ICE24 obtendrá trazabilidad operativa y evidencia suficiente para decidir la siguiente expansión.")
    doc.add_heading("5.1 Capacidad de referencia", level=2)
    add_table(doc, ["Elemento", "Supuesto de piloto"], [
        ("Cuentas", "1"), ("Sucursales", "Hasta 5"), ("Máquinas", "Hasta 25"), ("Usuarios", "Hasta 20"),
        ("Plataformas", "Escritorio web, Android PWA e iOS PWA"), ("Datos", "Sintéticos en desarrollo; reales sólo después de gates de producción")
    ], [3000, 6360])
    doc.add_heading("5.2 Métricas de los primeros 60 días", level=2)
    add_table(doc, ["Métrica", "Objetivo"], [
        ("Invitados que completan primer acceso", ">= 80%"),
        ("Órdenes asignadas completadas en PWA", ">= 70%"),
        ("Actividades completadas antes de vencer", ">= 85%"),
        ("Sincronizaciones sin soporte", ">= 98%"),
        ("Incidentes P0/P1 por pérdida/exposición", "0"),
        ("Usuarios piloto activos semanalmente", ">= 60%"),
    ], [7100, 2260])
    heading = doc.add_heading("5.3 Experiencia objetivo", level=2)
    heading.paragraph_format.page_break_before = True
    p(doc, "La interfaz será sobria, legible y orientada a la tarea: navegación lateral en escritorio, encabezados claros, tarjetas blancas sobre superficies azul pálido, estados semánticos visibles y acciones primarias en azul. En móvil se priorizan controles táctiles, cámara, avance local y estado de sincronización.")
    add_table(doc, ["Token", "Uso"], [
        ("Azul marino #04144C", "Marca, navegación, encabezados de alto énfasis"),
        ("Azul acción #0455C4", "Botones, enlaces, selección y foco"),
        ("Celeste #9CD5EF", "Acentos y superficies de apoyo"),
        ("Blanco", "Tarjetas y contenido principal"),
        ("Semánticos", "Éxito #16794D · advertencia #A15C00 · crítico #B42318 · offline #007B83"),
    ], [2700, 6660])

    page_break(doc)
    doc.add_heading("6. Así se verá la primera versión", level=1)
    p(doc, "Las siguientes pantallas son una visualización conceptual de la primera versión; no son una interfaz implementada ni sustituyen pruebas con usuarios. Traducen la arquitectura de información, la identidad visual y los flujos definidos en la documentación actual.")
    doc.add_heading("6.1 Escritorio: tablero de mantenimiento", level=2)
    add_picture(doc, desktop, width=Inches(6.35), alt="Concepto de tablero de mantenimiento de ICE24 OS en escritorio con navegación lateral, métricas y órdenes prioritarias.")
    add_caption(doc, "Figura 1. Concepto de la vista inicial para propietario u operaciones.")
    for item in [
        "Resumen inmediato de carga, vencimientos y cumplimiento.",
        "Filtro por cuenta y sucursal sin mezclar ámbitos.",
        "Acceso visible a órdenes, máquinas, historial y auditoría según permisos.",
        "Estados comprensibles por texto y color; el color nunca es la única señal.",
    ]:
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("6.2 Móvil: ejecución de una orden", level=2)
    add_picture(doc, mobile, width=Inches(3.45), alt="Concepto móvil de una orden de mantenimiento con checklist, estado sin conexión y guardado local.")
    add_caption(doc, "Figura 2. Concepto de la experiencia de campo en PWA instalada.")
    add_callout(doc, "Comportamiento clave", "La pérdida de red no borra el trabajo ni simula una sincronización. La aplicación muestra que está offline, guarda comandos y fotografías en el dispositivo, conserva identificadores idempotentes y sincroniza cuando vuelve la conectividad.", OFFLINE, "E6F6F6")
    heading = doc.add_heading("6.3 Pantallas mínimas", level=2)
    heading.paragraph_format.page_break_before = True
    add_table(doc, ["Superficie", "Pantallas del primer corte"], [
        ("Acceso", "Inicio de sesión, 2FA cuando aplique, recuperación y selección de cuenta"),
        ("Administración", "Resumen, sucursales, máquinas, usuarios/asociaciones"),
        ("Mantenimiento", "Listado, detalle, creación/asignación, ejecución, conflicto e historial"),
        ("Auditoría", "Eventos por recurso, actor, fecha y acción; exportación sólo autorizada"),
        ("Sistema", "Estado offline/sincronización, notificaciones y sesión"),
    ], [2200, 7160])

    page_break(doc)
    doc.add_heading("7. Escenario operativo de punta a punta", level=1)
    add_picture(doc, flow, width=Inches(6.35), alt="Flujo de siete pasos desde acceso y selección de contexto hasta cierre, historial y auditoría.")
    add_caption(doc, "Figura 3. Flujo principal que debe quedar demostrable en el MVP-1.")
    steps = [
        "El usuario inicia sesión mediante Keycloak. Los roles críticos usan TOTP y el navegador sólo recibe una cookie BFF segura.",
        "Selecciona la cuenta y sucursal. El servidor valida que la relación continúe activa.",
        "Un propietario u operador autorizado crea la orden, selecciona máquina, prioridad, fecha y técnico.",
        "El técnico abre su lista asignada, descarga la tarea y ejecuta checklist, notas y fotografías.",
        "Si no hay red, la PWA conserva comandos pendientes y muestra límites de almacenamiento y estado local.",
        "Al reconectar, sincroniza de forma idempotente. Un conflicto no se sobreescribe silenciosamente; se resuelve con contexto.",
        "La orden cerrada aparece en historial con evidencia, responsables, tiempos y eventos de auditoría.",
    ]
    add_numbered_list(doc, steps)
    doc.add_heading("7.1 Resultado demostrable", level=2)
    p(doc, "Durante una demostración, el equipo deberá poder alternar entre escritorio y móvil, provocar una pérdida de red controlada, completar al menos una actividad offline, reconectar y comprobar que el historial del servidor conserva exactamente una operación y la evidencia correspondiente.")

    page_break(doc)
    doc.add_heading("8. Arquitectura de solución en Google Cloud", level=1)
    add_picture(doc, architecture, width=Inches(6.35), alt="Diagrama de arquitectura de ICE24 OS en Google Cloud con Cloud Run, Cloud SQL, Storage, Keycloak, mensajería y observabilidad.")
    add_caption(doc, "Figura 4. Arquitectura objetivo del piloto; la topología final requiere validación de cuotas y costos regionales.")
    heading = doc.add_heading("8.1 Despliegues y servicios", level=2)
    heading.paragraph_format.page_break_before = True
    add_table(doc, ["Componente", "Servicio propuesto", "Responsabilidad"], [
        ("PWA + BFF", "Cloud Run", "Interfaz privada, sesión segura y llamadas de usuario"),
        ("API", "Cloud Run", "Reglas de negocio, autorización, persistencia y auditoría"),
        ("Portal público", "Cloud Run + CDN", "Sólo contenido deliberadamente publicado; posterior a MVP-1"),
        ("Identidad", "Keycloak en Cloud Run", "OIDC, credenciales, TOTP y sesiones globales"),
        ("Datos", "Cloud SQL PostgreSQL/PostGIS", "Datos transaccionales, geoespaciales y PITR"),
        ("Objetos", "Cloud Storage", "Cuarentena, originales y derivados con URLs temporales"),
        ("Asíncrono", "Pub/Sub, Cloud Tasks, Scheduler", "Notificaciones, trabajos, reintentos y programación"),
        ("Operación", "Artifact Registry, Terraform, OTel", "Promoción, infraestructura y telemetría"),
    ], [1900, 2600, 4860])

    page_break(doc)
    doc.add_heading("9. Stack y estructura técnica", level=1)
    add_table(doc, ["Capa", "Tecnología"], [
        ("Runtime/lenguaje", "Node.js 24 LTS · TypeScript estricto"),
        ("Monorepo", "pnpm workspaces · Turborepo"),
        ("Web/PWA", "Next.js · React · App Router · Dexie/IndexedDB"),
        ("API/workers", "NestJS · REST/OpenAPI"),
        ("Persistencia", "PostgreSQL/PostGIS · Prisma + SQL explícito"),
        ("Pruebas", "Vitest unitario · Playwright E2E/PDF · Testcontainers"),
        ("Entrega", "Docker · Terraform · Artifact Registry · CI/CD"),
        ("Observabilidad", "OpenTelemetry · Cloud Logging/Monitoring/Trace"),
    ], [2500, 6860])
    doc.add_heading("9.1 Arquitectura de aplicación", level=2)
    p(doc, "Se adopta un monolito modular: un contrato de API independiente, módulos de dominio con límites explícitos y procesamiento asíncrono cuando aporta confiabilidad. No se justifican microservicios ni Kubernetes en el piloto.")
    doc.add_heading("9.2 Estructura de repositorio prevista", level=2)
    for item in [
        "apps/web — PWA privada y BFF.",
        "apps/api — API NestJS y contrato OpenAPI.",
        "apps/portal — portal público futuro.",
        "apps/worker y apps/pdf — procesamiento asíncrono y documentos.",
        "packages/* — dominio, autorización, UI, contratos, configuración y utilidades compartidas.",
        "infra/* — Terraform, entornos y políticas de despliegue.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("9.3 Política de versiones", level=2)
    p(doc, "Fase 1 fijará versiones exactas en packageManager, runtime, imágenes y lockfile. No se usarán RC/beta/canary en componentes críticos. Parches críticos se aplicarán dentro de 72 horas cuando exista una corrección compatible.")

    page_break(doc)
    doc.add_heading("10. Datos, seguridad e identidad", level=1)
    doc.add_heading("10.1 Identidad y sesión", level=2)
    add_table(doc, ["Control", "Valor propuesto"], [
        ("Protocolo", "OIDC Authorization Code + PKCE mediante BFF"),
        ("Access token", "5 minutos"),
        ("Sesión inactiva", "30 min administración; hasta 8 h en campo con bloqueo/reauth sensible"),
        ("Duración absoluta", "12 horas"),
        ("Recordar dispositivo", "Deshabilitado en piloto"),
        ("2FA", "TOTP obligatorio para ICE24 Admin, propietarios, responsables sanitarios y publicadores"),
        ("Recuperación", "Enlace de un uso; casos excepcionales con dos personas y auditoría"),
    ], [2800, 6560])
    doc.add_heading("10.2 Invariantes de datos", level=2)
    for item in [
        "Todas las entidades de negocio quedan vinculadas al tenant/cuenta aplicable.",
        "Los identificadores de cliente para sincronización permiten idempotencia y detección de duplicados.",
        "Los objetos binarios no se almacenan en PostgreSQL; sólo metadatos y referencias.",
        "Los originales privados no se convierten en públicos por compartir un enlace interno.",
        "Auditoría append-only para acciones relevantes; datos restringidos no aparecen en logs.",
        "PostGIS y las reglas geoespaciales permanecen bajo SQL explícito cuando Prisma no alcanza.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("10.3 Defensa en profundidad", level=2)
    p(doc, "Cookie HttpOnly/Secure/SameSite, CSP y cabeceras, validación de entrada, autorización por comando, secretos en Secret Manager, cifrado administrado/KMS, red privada para base de datos, escaneo antimalware, mínimo privilegio de servicios y pruebas negativas por cuenta, acción, ámbito y sensibilidad.")

    page_break(doc)
    doc.add_heading("11. Offline, archivos, observabilidad y continuidad", level=1)
    doc.add_heading("11.1 Modo offline", level=2)
    add_table(doc, ["Condición", "Respuesta esperada"], [
        ("Red estable", "Operación completa"),
        ("Alta latencia", "Feedback visible y cargas reanudables"),
        ("Pérdida intermitente", "Reintentos idempotentes y estado visible"),
        ("Sin red", "Sólo tareas de mantenimiento previamente sincronizadas"),
        ("Límite local", "7 días, 50 tareas, 200 fotos o 1 GB; advertencia al 80%"),
    ], [2300, 7060])
    doc.add_heading("11.2 Ciclo de archivos", level=2)
    file_steps = [
        "Carga a bucket privado de cuarentena mediante autorización temporal.",
        "ClamAV en Cloud Run, activado por Eventarc; fallo o timeout conserva cuarentena.",
        "Sólo un resultado limpio permite promoción a originales/derivados.",
        "Descargas con URLs firmadas de corta duración; exportaciones generadas por tiempo limitado.",
    ]
    add_numbered_list(doc, file_steps)
    doc.add_heading("11.3 Operación y continuidad", level=2)
    add_table(doc, ["Indicador", "Objetivo"], [
        ("Disponibilidad privada", "99.5% mensual"),
        ("Lectura API común", "p95 <= 500 ms"),
        ("Escritura API común", "p95 <= 800 ms"),
        ("LCP principal", "p75 <= 2.5 s"),
        ("Datos transaccionales", "RPO 15 min · RTO 4 h"),
        ("Restauración", "Prueba trimestral antes de disponibilidad general"),
    ], [3500, 5860])

    page_break(doc)
    doc.add_heading("12. Entornos, costos y despliegue", level=1)
    add_table(doc, ["Entorno", "Uso", "Datos y aislamiento"], [
        ("Local", "Desarrollo", "Sintéticos · Docker Compose"),
        ("CI", "Pruebas efímeras", "Fábricas/sintéticos · Testcontainers"),
        ("Development", "Integración accesible por URL", "Sintéticos · proyecto GCP no productivo"),
        ("Staging", "Pruebas de release", "Anonimizados/sintéticos · recursos separados"),
        ("Production", "Piloto y operación", "Reales · proyecto productivo, mínimo privilegio y Cloud SQL HA"),
    ], [1500, 2700, 5160])
    doc.add_heading("12.1 Estrategia de despliegue", level=2)
    for item in [
        "Terraform es la vía normal de creación de infraestructura.",
        "CI construye una imagen inmutable y promueve la misma imagen entre entornos.",
        "Development y staging escalan a cero donde sea seguro; Keycloak mantiene una instancia mínima.",
        "Producción y no producción viven en proyectos GCP distintos.",
        "Los dominios app, api, auth, public y files quedan sujetos a compra y configuración DNS.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("12.2 Presupuesto de control", level=2)
    add_callout(doc, "Objetivo del piloto", "Mantener la plataforma entre USD 100 y 300 mensuales. Tope provisional: USD 400/mes en producción y USD 150/mes combinados para development/staging, excluyendo impuestos, soporte empresarial y crecimiento extraordinario.", WARNING, "FFF6E8")
    p(doc, "Antes de aprovisionar deben ejecutarse la calculadora de precios de Google Cloud y presupuestos con alertas al 50%, 80% y 100%. Cloud SQL y la instancia mínima de Keycloak son los principales costos persistentes.")

    page_break(doc)
    doc.add_heading("13. Calidad y criterios de aceptación", level=1)
    doc.add_heading("13.1 Pirámide de pruebas", level=2)
    add_table(doc, ["Nivel", "Herramienta", "Cobertura mínima"], [
        ("Unidad", "Vitest", "Dominio, políticas, validaciones y adaptadores"),
        ("Integración", "Vitest + Testcontainers", "PostgreSQL/PostGIS, repositorios, migraciones y colas"),
        ("Contrato", "OpenAPI", "Compatibilidad web/API y errores estables"),
        ("E2E", "Playwright", "Acceso, orden, offline, sincronización, conflicto y auditoría"),
        ("Seguridad", "Automatizada + revisión", "Autorización negativa, sesiones, archivos y secretos"),
        ("Accesibilidad", "Automatizada + manual", "WCAG 2.2 AA en flujos MVP-1"),
    ], [1700, 2500, 5160])
    doc.add_heading("13.2 Definition of Done del primer flujo", level=2)
    for item in [
        "Historia trazada a requisito y criterio de aceptación.",
        "Pruebas positivas y negativas por acción, cuenta, ámbito y sensibilidad.",
        "Telemetría útil sin datos restringidos.",
        "Manejo de error, reintento e idempotencia demostrado.",
        "Accesibilidad por teclado, lector, contraste y tamaño táctil.",
        "Documentación y runbook actualizados; rollback probado cuando aplique.",
        "Despliegue en development mediante pipeline, sin secretos en Git.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("13.3 Compatibilidad", level=2)
    p(doc, "Nivel A: Chrome y Edge actuales en Windows 11, Safari actual en macOS, Chrome Android 11+ y PWA en las dos versiones mayores actuales de iOS. Se difieren modo oscuro y densidad configurable; no se difieren contraste, foco, zoom ni tamaño táctil.")

    page_break(doc)
    doc.add_heading("14. Estado de Fase 0 y gates", level=1)
    status_rows = [
        ("F0-01", "RACI", "En revisión", "Nombres y aceptación"),
        ("F0-02", "MVP web", "En revisión", "Dirección/Producto"),
        ("F0-03", "Decisiones", "En curso avanzado", "Resolver P0/P1"),
        ("F0-04", "Plataforma", "En revisión", "Presupuesto/privacidad"),
        ("F0-05", "SLO/continuidad", "En revisión", "Operación/Seguridad/Legal"),
        ("F0-06", "Stack/Vitest", "Lista para adopción", "Lockfile en Fase 1"),
        ("F0-07", "Identidad", "En revisión", "PoC Keycloak/recuperación"),
        ("F0-08", "Autorización", "En revisión", "Validación por dominio"),
        ("F0-09", "Códigos/folios", "En revisión", "Spike en Fase 1"),
        ("F0-10", "Formatos Excel", "Bloqueada", "3 archivos reales anonimizados"),
        ("F0-11", "Plantillas", "Bloqueada", "Manuales/plantillas aprobadas"),
        ("F0-12", "Sanidad/jurídico", "Bloqueada", "Dictamen de responsables"),
        ("F0-13", "Proveedores", "En revisión", "Cuentas/cuotas/presupuesto"),
        ("F0-14", "Soporte", "En revisión", "Dispositivos del piloto"),
        ("F0-15", "Incidentes", "En revisión", "Guardia/canales reales"),
        ("F0-16", "Backlog", "Terminado documentalmente", "Orden de releases"),
    ]
    add_table(doc, ["Tarea", "Entregable", "Estado", "Gate"], status_rows, [1100, 2300, 2550, 3410])
    add_callout(doc, "Próximo gate recomendado", "Nombrar la RACI, aprobar el MVP-1, autorizar presupuesto/privacidad y ejecutar Fase 1 con un esqueleto sin datos sensibles. En paralelo, recopilar Excel, manuales y dictámenes externos.", SUCCESS, "EAF6F0")

    page_break(doc)
    doc.add_heading("15. Riesgos, bloqueos y decisiones abiertas", level=1)
    add_table(doc, ["Riesgo o bloqueo", "Impacto", "Tratamiento recomendado"], [
        ("Excel reales no disponibles", "No se validan layouts ni reglas de importación", "Recibir 3 muestras anonimizadas y construir fixtures versionados"),
        ("Plantillas aprobadas ausentes", "No se puede automatizar documentación regulada", "Catálogo y responsables por plantilla antes de MVP-2"),
        ("Dictamen sanitario/jurídico pendiente", "No publicar ni fijar retenciones definitivas", "Taller y aprobación registrada"),
        ("Presupuesto/privacidad GCP", "Producción no autorizada", "Calculadora, DPA, residencia y alertas"),
        ("Keycloak en Cloud Run", "Costo fijo y componente crítico", "PoC de sesiones, backup, recuperación y actualización"),
        ("Conectividad/dispositivo real", "Offline podría fallar en campo", "Inventario de hardware y pruebas con red degradada"),
        ("Proveedor correo/mapas", "Costo y transferencia de datos", "Adaptadores, cuotas, DPA y minimización"),
        ("Expansión prematura", "Retraso del valor principal", "Proteger alcance MVP-1 con gates por evidencia"),
    ], [3150, 2800, 3410])
    doc.add_heading("15.1 Decisiones que deben cerrarse primero", level=2)
    decision_steps = [
        "Personas responsables en la RACI y dueños de guardia/recuperación.",
        "Aprobación formal del alcance, métricas y cuenta piloto.",
        "Presupuesto inicial, organización/facturación GCP y dominios.",
        "PoC de Keycloak y política de sesión de campo.",
        "Clasificación/retención aprobada y reglas de privacidad.",
    ]
    add_numbered_list(doc, decision_steps)

    page_break(doc)
    doc.add_heading("16. Ruta recomendada", level=1)
    p(doc, "La mejor ruta para llegar a una aplicación web funcional es mantener el foco en la rebanada vertical de mantenimiento y desplegar continuamente desde Fase 1.")
    add_table(doc, ["Paso", "Resultado", "Gate de salida"], [
        ("1. Gobierno", "RACI, MVP y presupuesto aprobados", "Responsables y autoridad explícita"),
        ("2. Fundación", "Monorepo, CI, IaC, URL development y observabilidad mínima", "Build/test/deploy reproducible"),
        ("3. Identidad", "Keycloak, BFF, contexto y autorización base", "Pruebas negativas y recuperación"),
        ("4. Núcleo", "Cuenta, sucursal, máquina y asociaciones", "Separación tenant demostrada"),
        ("5. Flujo MVP", "Orden, checklist, fotos, offline, sincronización e historial", "E2E completo en staging"),
        ("6. Piloto", "Producción, usuarios reales, soporte y métricas", "Gates de seguridad/privacidad cerrados"),
        ("7. Expansión", "MVP-2 técnico-sanitario", "F0-10, F0-11 y F0-12 cerradas"),
    ], [1400, 5100, 2860])
    doc.add_heading("16.1 Primer incremento de implementación", level=2)
    for item in [
        "Crear apps/web y apps/api con endpoint de salud y página de acceso provisional.",
        "Configurar Vitest, Playwright, lint, typecheck y pipeline de CI.",
        "Provisionar development mínimo por Terraform y desplegar por imagen inmutable.",
        "Incorporar contrato OpenAPI, telemetría y manejo de errores desde el primer endpoint.",
        "Usar sólo datos sintéticos hasta que el gate de producción esté aprobado.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Criterio de avance", "Una URL accesible no es el MVP terminado; es el primer peldaño. La primera versión se considera funcional cuando el flujo de mantenimiento completo —incluido un tramo offline— se demuestra, audita y puede operar con seguridad en la cuenta piloto.")

    page_break(doc)
    doc.add_heading("Anexo A. Registro de decisiones vigentes", level=1)
    add_table(doc, ["ADR", "Decisión", "Estado"], [
        ("ADR-015", "Google Cloud, región México y entornos", "Línea base; presupuesto/privacidad pendientes"),
        ("ADR-016", "Stack, Node 24 LTS y Vitest", "Aceptada para Fase 1"),
        ("ADR-017", "Keycloak, BFF, 2FA, recuperación y sesiones", "En revisión"),
        ("ADR-018", "SLO, continuidad y retención provisional", "En revisión"),
        ("ADR-019", "Storage, ClamAV, Resend y Mapbox", "En revisión de costo/privacidad"),
        ("ADR-020", "Identificadores y folios", "En revisión; validación en Fase 1"),
    ], [1500, 5200, 2660])
    doc.add_heading("Anexo B. Documentos fuente", level=1)
    p(doc, "Este informe sintetiza los archivos versionados del repositorio; ante una diferencia, prevalece la fuente específica más reciente y la decisión aprobada por su responsable.")
    sources = [
        "context/ICE24_OS_Documento_Maestro_Requerimientos_v1.0.md",
        "context/ICE24_OS_PRD_v1.0.md",
        "context/ICE24_OS_TRD_v1.0.md",
        "context/Architecture.md · Database.md · API.md · AppFlow.md · UI_UX.md",
        "context/Implementation_Plan.md · PROJECT_RULES.md · TASKS.md",
        "docs/product/* · docs/decisions/* · docs/backlog/* · docs/tasks/*",
    ]
    for source in sources:
        add_bullet(doc, source)
    doc.add_heading("Anexo C. Notas de interpretación", level=1)
    for item in [
        "Las cifras de costo son límites de control, no cotizaciones comerciales.",
        "La capacidad del piloto no es un límite comercial codificado.",
        "Las pantallas son conceptos de diseño; el detalle final se valida con prototipo y usuarios.",
        "Los valores de retención permanecen provisionales hasta dictamen legal y sanitario.",
        "La fecha de estado de Fase 0 en su fuente es 17/08/2026; este documento tiene corte al 18/08/2026.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
